import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(BASE_DIR, "dokumen", "LAPORAN_DATASET.docx")

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_heading(p, text, level=1):
    p.text = text
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.runs[0]
    run.font.name = 'Arial'
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 51, 51)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title_p.add_run('LAPORAN EKSPLORASI & PENGELOLAAN DATASET\n')
t_run.font.name = 'Arial'
t_run.font.size = Pt(15)
t_run.bold = True
t_run.font.color.rgb = RGBColor(0, 51, 102)

sub_run = title_p.add_run('Klasifikasi Penyakit Daun Bawang Merah Berbasis CNN (MobileNetV2)')
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 1
p1 = doc.add_paragraph()
style_heading(p1, '1. Ringkasan Eksekutif Dataset', 1)
doc.add_paragraph(
    'Dataset yang digunakan dalam penelitian ini difokuskan pada identifikasi kondisi kesehatan '
    'tanaman bawang merah (Allium cepa var. aggregatum) berdasarkan citra visual daun. '
    'Dataset terdiri atas 4 kelas klasifikasi, yaitu 3 kelas penyakit/kondisi tanaman bawang merah '
    'dan 1 kelas negatif (bukan daun bawang merah) untuk mencegah false positive pada lingkungan implementasi riil.'
)

# Table 1: Summary
t1 = doc.add_table(rows=6, cols=2)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
t1_data = [
    ('Total Jumlah Sampel', '2.342 Citra'),
    ('Jumlah Kelas Target', '4 Kelas (moler, non_bawang, sehat, trotol)'),
    ('Integritas Berkas', '100% Valid (0 File Corrupt)'),
    ('Rasio Pembagian Dataset', '70% Training, 15% Validation, 15% Testing'),
    ('Dimensi Input Model', '224 x 224 piksel (3 Channel RGB)'),
    ('Normalisasi Nilai Piksel', 'Rescale 1/255 -> [0.0, 1.0]')
]
for i, (k, v) in enumerate(t1_data):
    row = t1.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(row.cells[0], 'F0F4F8')
    if i % 2 == 1:
        set_cell_shading(row.cells[1], 'FAFAFA')

# Section 2
p2 = doc.add_paragraph()
style_heading(p2, '2. Karakteristik Visual Setiap Kelas', 1)
doc.add_paragraph(
    'Setiap kelas memiliki ciri visual dan peran spesifik dalam pelatihan:\n'
    '• Moler (Layu Fusarium / Fusarium oxysporum): Daun melintir abnormal, layu terkulai, spiral/keriting, menguning secara bertahap.\n'
    '• Trotol (Bercak Ungu / Alternaria porri): Bercak melekuk putih kelabu dengan tepi keunguan, berkembang menjadi lesi konsentris gelap.\n'
    '• Sehat: Daun hijau segar, silindris tegak, mulus tanpa bercak atau pembelitan.\n'
    '• Non Bawang (Out-of-Distribution): Objek non-bawang seperti rumput liar, tanah, tangan, dan tanaman lain guna mencegah false positive.'
)

# Section 3
p3 = doc.add_paragraph()
style_heading(p3, '3. Statistik Sampel & Format Berkas', 1)
t2 = doc.add_table(rows=5, cols=5)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Kelas', 'Jumlah Citra', 'Persentase', 'Format File', 'Resolusi Asli Terbanyak']
for j, h in enumerate(headers):
    t2.rows[0].cells[j].text = h
    t2.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(t2.rows[0].cells[j], '2B5C8F')
    t2.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

t2_data = [
    ('moler', '500', '21.35%', 'JPEG (162), PNG (338)', '500x500, 4000x3000, 3060x4080'),
    ('non_bawang', '842', '35.95%', 'JPEG (829), WEBP (10), PNG (3)', '224x224, 626x418, 1280x720'),
    ('sehat', '500', '21.35%', 'JPEG (500)', '3060x3060, 640x640, 3060x4080'),
    ('trotol', '500', '21.35%', 'JPEG (500)', '640x640, 3060x4080, 256x256')
]
for i, row_data in enumerate(t2_data):
    row = t2.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if i % 2 == 1:
            set_cell_shading(row.cells[j], 'F5F7FA')

# Section 4
p4 = doc.add_paragraph()
style_heading(p4, '4. Pembagian Dataset (Data Splitting 70:15:15)', 1)
t3 = doc.add_table(rows=6, cols=5)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['Kelas Target', 'Train (70%)', 'Val (15%)', 'Test (15%)', 'Total Citra']
for j, h in enumerate(headers3):
    t3.rows[0].cells[j].text = h
    t3.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(t3.rows[0].cells[j], '2B5C8F')
    t3.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

t3_data = [
    ('moler', '350', '75', '75', '500'),
    ('non_bawang', '589', '126', '127', '842'),
    ('sehat', '350', '75', '75', '500'),
    ('trotol', '350', '75', '75', '500'),
    ('TOTAL', '1.639', '351', '352', '2.342')
]
for i, row_data in enumerate(t3_data):
    row = t3.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if i == len(t3_data) - 1:
            row.cells[j].paragraphs[0].runs[0].font.bold = True
            set_cell_shading(row.cells[j], 'DCE6F1')
        elif i % 2 == 1:
            set_cell_shading(row.cells[j], 'F5F7FA')

# Section 5
p5 = doc.add_paragraph()
style_heading(p5, '5. Konfigurasi Pra-Pemrosesan & Augmentasi Data', 1)
doc.add_paragraph(
    'Pra-pemrosesan meliputi resizing citra menjadi 224x224 piksel dan normalisasi rescale 1/255. '
    'Pada subset data latih (training), diaplikasikan augmentasi real-time untuk mencegah overfitting:\n'
    '• Rotation Range: Rotasi acak hingga ±25°\n'
    '• Width & Height Shift: Pergeseran posisi 20%\n'
    '• Shear Range: Transformasi sudut 20%\n'
    '• Zoom Range: Perbesaran/perkecilan acak 30%\n'
    '• Brightness Range: Variasi pencahayaan 70% - 130%\n'
    '• Channel Shift: Variasi kanal warna RGB 20.0\n'
    '• Flip: Horizontal dan Vertical Flip (True)\n'
    '• Fill Mode: nearest'
)

# Save document
doc.save(DOCX_PATH)
print('Berhasil membuat file:', DOCX_PATH)
